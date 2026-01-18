"""Lightweight, non-LLM document analysis for project workspaces.

Supported formats: docx, pdf, txt, md. Outputs per-file summaries with hashes
for deduplication plus simple heuristics for roles, dates, metrics, and skills."""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:
    PdfReader = None


SUPPORTED_DOC_EXTS = {".docx", ".pdf", ".txt", ".md"}


def compute_sha256(path: Path) -> str:
    """
    Compute a SHA256 hash of a file for deduplication.

    Args:
        path (Path): File path to hash.

    Returns:
        str: Hex digest of the file content.
    """
    digest = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            digest.update(chunk)
    return digest.hexdigest()


@dataclass
class ParsedDoc:
    text: str
    headings: List[str] = field(default_factory=list)


class DocumentAnalyzer:
    """
    Analyze text-bearing project documents without external services.

    Provides:
      - Content hashing for deduplication
      - Shallow text extraction per format
      - Heuristic extraction of metrics, dates, roles, and skills
    """

    def __init__(
        self,
        root: Path,
        files: Optional[Iterable[Path]] = None,
        known_hashes: Optional[Dict[str, str]] = None,
    ):
        """
        Initialize the analyzer with a project root and optional prior hashes.
        Args:
            root (Path): Project root to scan for documents.
            files (Optional[Iterable[Path]]): Optional iterable of files to analyze.
            known_hashes (Optional[Dict[str, str]]): Existing hash→path map to flag duplicates.
        Returns:
            None
        """
        self.root = Path(root)
        self.files = list(files) if files is not None else None
        self.known_hashes: Dict[str, str] = dict(known_hashes or {})

    def analyze(self) -> Dict[str, Any]:
        """
        Analyze supported documents under the root path and return structured findings.
        Args:
            None
        Returns:
            Dict[str, Any]: Contains per-document details, duplicates, hash index,
                summary stats, and any errors encountered.
        """
        documents: List[Dict[str, Any]] = []
        duplicates: List[Dict[str, str]] = []
        errors: List[str] = []

        if self.files is None and not self.root.exists():
            return {
                "documents": [],
                "duplicates": [],
                "summary": {"unique_documents": 0, "duplicate_documents": 0, "total_words": 0, "by_format": {}},
                "hash_index": {},
                "errors": [f"Root path not found: {self.root}"],
            }

        paths = self.files if self.files is not None else sorted(self.root.rglob("*"))

        for path in paths:
            if not path.is_file():
                continue

            if path.name.startswith("._") or "__MACOSX" in path.parts:
                continue

            suffix = path.suffix.lower()
            if suffix not in SUPPORTED_DOC_EXTS:
                continue

            try:
                rel_path = str(path.relative_to(self.root))
            except ValueError:
                rel_path = str(path)

            try:
                file_hash = compute_sha256(path)
            except Exception as e:
                errors.append(f"hash_failed:{rel_path}:{e}")
                continue

            if file_hash in self.known_hashes:
                duplicates.append(
                    {"path": rel_path, "hash": file_hash, "duplicate_of": self.known_hashes[file_hash]}
                )
                continue

            try:
                parsed = self._extract_content(path, suffix)
            except Exception as e:
                errors.append(f"parse_failed:{rel_path}:{e}")
                continue

            record = self._build_record(rel_path, file_hash, suffix, parsed)
            documents.append(record)
            self.known_hashes[file_hash] = rel_path

        summary = self._build_summary(documents, duplicates)

        return {
            "documents": documents,
            "duplicates": duplicates,
            "summary": summary,
            "hash_index": self.known_hashes,
            "errors": errors,
        }

    def _build_summary(self, documents: List[Dict[str, Any]], duplicates: List[Dict[str, str]]) -> Dict[str, Any]:
        """
        Summarize document counts and word totals for the analysis.
        Args:
            documents (List[Dict[str, Any]]): Unique document records.
            duplicates (List[Dict[str, str]]): Duplicate document entries.
        Returns:
            Dict[str, Any]: Aggregate counts for unique/duplicate docs, words, and formats.
        """
        total_words = sum(doc.get("word_count", 0) for doc in documents)
        by_format: Dict[str, int] = {}
        for doc in documents:
            fmt = doc.get("format")
            if fmt:
                by_format[fmt] = by_format.get(fmt, 0) + 1

        return {
            "unique_documents": len(documents),
            "duplicate_documents": len(duplicates),
            "total_words": total_words,
            "by_format": by_format,
        }

    def _build_record(
        self,
        rel_path: str,
        file_hash: str,
        suffix: str,
        parsed: ParsedDoc,
    ) -> Dict[str, Any]:
        """
        Build the per-document record with extracted signals.
        Args:
            rel_path (str): File path relative to root.
            file_hash (str): SHA256 fingerprint of the file.
            suffix (str): File extension.
            parsed (ParsedDoc): Extracted text and headings.
        Returns:
            Dict[str, Any]: Structured document payload with metadata and heuristics.
        """
        text = parsed.text.strip()
        preview = " ".join(text.split())[:400]
        headings = parsed.headings[:10]
        lower_text = text.lower()

        return {
            "path": rel_path,
            "format": suffix.lstrip(".").upper(),
            "sha256": file_hash,
            "word_count": self._word_count(text),
            "headings": headings,
            "preview": preview,
            "metrics": self._extract_metrics(text),
            "dates": self._extract_dates(text),
            "roles": self._extract_roles(lower_text, headings),
            "skills": self._extract_skills(lower_text),
        }

    def _extract_content(self, path: Path, suffix: str) -> ParsedDoc:
        """
        Route file reading based on extension.
        Args:
            path (Path): File path to read.
            suffix (str): Lowercased file extension.
        Returns:
            ParsedDoc: Extracted text and headings.
        """
        suffix = suffix.lower()
        if suffix == ".docx":
            return self._read_docx(path)
        if suffix == ".pdf":
            return self._read_pdf(path)
        if suffix == ".md":
            return self._read_markdown(path)
        return self._read_plain_text(path)

    def _read_docx(self, path: Path) -> ParsedDoc:
        """
        Read DOCX content and collect paragraphs, headings, and table text.
        Args:
            path (Path): DOCX file path.
        Returns:
            ParsedDoc: Combined text and detected headings.
        """
        if Document is None:
            raise ImportError("python-docx is not installed")
        doc = Document(path)
        paragraphs = []
        headings = []

        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
            paragraphs.append(text)
            try:
                style_name = (p.style.name or "").lower()
                if "heading" in style_name:
                    headings.append(text)
            except Exception:
                pass

        # Capture table text to preserve metrics
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return ParsedDoc(text="\n".join(paragraphs), headings=headings)

    def _read_pdf(self, path: Path) -> ParsedDoc:
        """
        Read PDF content using pypdf and join page text.
        Args:
            path (Path): PDF file path.
        Returns:
            ParsedDoc: Combined text from all pages.
        """
        if PdfReader is None:
            raise ImportError("pypdf is not installed")
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            try:
                txt = page.extract_text() or ""
            except Exception:
                txt = ""
            if txt:
                pages.append(txt)
        return ParsedDoc(text="\n".join(pages))

    def _read_plain_text(self, path: Path) -> ParsedDoc:
        """
        Read a plain text file with utf-8 fallback to latin-1 on decode errors.
        Args:
            path (Path): Text file path.
        Returns:
            ParsedDoc: Raw text content.
        """
        try:
            text = path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = path.read_text(encoding="latin-1")
        return ParsedDoc(text=text)

    def _read_markdown(self, path: Path) -> ParsedDoc:
        """
        Read markdown files and capture heading lines.
        Args:
            path (Path): Markdown file path.
        Returns:
            ParsedDoc: Text content with detected headings.
        """
        parsed = self._read_plain_text(path)
        headings = []
        for line in parsed.text.splitlines():
            match = re.match(r"^#{1,6}\s+(.*)$", line.strip())
            if match:
                headings.append(match.group(1).strip())
        parsed.headings = headings
        return parsed

    def _word_count(self, text: str) -> int:
        """
        Count words in a block of text.
        Args:
            text (str): Text to count.
        Returns:
            int: Number of word tokens detected.
        """
        return len(re.findall(r"\b\w+\b", text))

    def _extract_metrics(self, text: str) -> List[str]:
        """
        Find metric-like patterns such as percentages and counts.
        Args:
            text (str): Text to search.
        Returns:
            List[str]: Unique metric strings found.
        """
        patterns = [
            r"\b\d{1,3}(?:,\d{3})*(?:\.\d+)?\s*%",
            r"\b\d+(?:\.\d+)?\s*(?:users|clients|tickets|tests|deployments|issues|prs|pull requests)\b",
            r"\b\d+(?:\.\d+)?\s*(?:x|X|times)\b",
        ]
        metrics: List[str] = []
        for pat in patterns:
            metrics.extend(re.findall(pat, text, flags=re.IGNORECASE))
        # Deduplicate while preserving order
        seen = set()
        unique_metrics = []
        for m in metrics:
            key = m.lower()
            if key in seen:
                continue
            seen.add(key)
            unique_metrics.append(m.strip())
        return unique_metrics[:10]

    def _extract_dates(self, text: str) -> List[str]:
        """
        Detect date-like ranges and years in text.
        Args:
            text (str): Text to search.
        Returns:
            List[str]: Unique date strings found.
        """
        # Capture ranges and single years/months
        patterns = [
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}\s*(?:-|to|through|–|—)\s*(?:Present|Now|\d{4})",
            r"\b\d{4}\s*(?:-|to|through|–|—)\s*(?:Present|Now|\d{4})",
            r"\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Sept|Oct|Nov|Dec)[a-z]*\s+\d{4}\b",
            r"\b(20\d{2}|19\d{2})\b",
        ]
        dates: List[str] = []
        for pat in patterns:
            dates.extend(re.findall(pat, text, flags=re.IGNORECASE))
        seen = set()
        unique_dates = []
        for d in dates:
            key = d.lower() if isinstance(d, str) else str(d).lower()
            if key in seen:
                continue
            seen.add(key)
            unique_dates.append(d if isinstance(d, str) else str(d))
        return unique_dates[:10]

    def _extract_roles(self, lower_text: str, headings: List[str]) -> List[str]:
        """
        Identify role-related keywords in text and headings.
        Args:
            lower_text (str): Lowercased document text.
            headings (List[str]): Extracted headings.
        Returns:
            List[str]: Sorted list of role titles found.
        """
        role_keywords = [
            "engineer",
            "developer",
            "manager",
            "lead",
            "researcher",
            "analyst",
            "designer",
            "architect",
            "consultant",
        ]
        roles = set()
        text_slice = lower_text[:1200]
        for keyword in role_keywords:
            if keyword in text_slice:
                roles.add(keyword.title())

        for heading in headings:
            h_lower = heading.lower()
            for keyword in role_keywords:
                if keyword in h_lower:
                    roles.add(keyword.title())
        return sorted(roles)

    def _extract_skills(self, lower_text: str) -> List[str]:
        """
        Identify predefined skills from document text.
        Args:
            lower_text (str): Lowercased document text.
        Returns:
            List[str]: Ordered unique skills detected.
        """
        skill_terms = {
            "python": "Python",
            "java": "Java",
            "javascript": "JavaScript",
            "typescript": "TypeScript",
            "c++": "C++",
            "c+": "C++",
            "cpp": "C++",
            "c#": "C#",
            "golang": "Go",
            "go language": "Go",
            "rust": "Rust",
            "sql": "SQL",
            "react": "React",
            "node.js": "Node.js",
            "nodejs": "Node.js",
            "express": "Express",
            "django": "Django",
            "flask": "Flask",
            "fastapi": "FastAPI",
            "spring": "Spring",
            "angular": "Angular",
            "vue": "Vue",
            "next.js": "Next.js",
            "nestjs": "NestJS",
            "pytorch": "PyTorch",
            "tensorflow": "TensorFlow",
            "docker": "Docker",
        }
        detected = []
        for needle, label in skill_terms.items():
            if needle in lower_text:
                detected.append(label)
        # Deduplicate while preserving order
        seen = set()
        unique_skills = []
        for s in detected:
            if s in seen:
                continue
            seen.add(s)
            unique_skills.append(s)
        return unique_skills
